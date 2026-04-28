# pyright: reportMissingImports=false, reportUnknownVariableType=false, reportUnknownMemberType=false, reportUnknownArgumentType=false, reportUnknownParameterType=false, reportAny=false, reportAttributeAccessIssue=false, reportUnusedCallResult=false, reportCallIssue=false, reportGeneralTypeIssues=false, reportArgumentType=false, reportOptionalSubscript=false, reportReturnType=false, reportMissingTypeArgument=false, reportDeprecated=false, reportUnannotatedClassAttribute=false
"""DOCX form-fill — replace ``{{placeholders}}`` in a .docx using corpus answers.

Pain point: Vendor questionnaires, DDQs, security forms arrive as Word files with
blanks. Ops copy-pastes from policy PDFs and tracks citations in a sidecar sheet.
This recipe: a ``.docx`` with ``{{field_name}}`` tokens → filled .docx + cited
evidence sidecar (JSON), every answer grounded in real ``[chunk:<uuid>]`` tags.

Framework: pydantic-ai + python-docx.
Tools used: list_contents, search_knowledge, read.
Output: <template-stem>.filled.docx  +  <template-stem>.citations.json
"""

import argparse
import asyncio
import json
import os
import re
import sys
from pathlib import Path

from docx import Document as DocxDocument
from pydantic import BaseModel, Field
from pydantic_ai import Agent
from pydantic_ai.mcp import MCPServerStdio

CORPUS_FOLDER = os.environ.get("POLICIES_FOLDER_ID", "")
PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-]+)\s*\}\}")


class Citation(BaseModel):
    chunk_id: str
    document_name: str
    snippet: str = Field(..., max_length=240)


class FieldAnswer(BaseModel):
    value: str = Field(..., max_length=1200)
    citations: list[Citation] = Field(..., min_length=1, max_length=4)


PROMPT = """Available MCP tools (use ONLY these): search_knowledge, search_keyword, read, find, list_contents, get_info. There is NO 'cite' tool, NO 'citation' tool, NO 'add_citation' tool. Build the citations field of your final response directly from the search_knowledge / read results — DO NOT try to call any tool to construct citations. You answer a single form field from corporate policy. Use MCP tools to ground
KS workflow (do NOT skip):
1. Ask Knowledge Stack specific natural-language questions about the input. Never use folder UUIDs or path_part_id filters in your queries.
2. search_knowledge returns hits. EACH hit is a JSON object with TWO distinct UUIDs: chunk_id (for citation only) and path_part_id (the chunk's path-tree node, used for read). The text field on the hit is empty.
3. To retrieve the chunk content, call read(path_part_id=<hit.path_part_id>). DO NOT pass chunk_id to read — read() returns 404 on chunk_ids. If you see a 404, you used the wrong UUID; switch to path_part_id from the SAME hit. The read() output ends in a [chunk:<uuid>] marker — that uuid is the citation.chunk_id.
4. Build every output field ONLY from chunk text you read. Never invent facts. If the corpus has nothing relevant, mark the field accordingly (e.g. confidence='low' or 'not in corpus — upload data to proceed').
5. Populate every citation with chunk_id (verbatim from the marker), document_name (filename from read() output's metadata or materialized_path), and snippet (verbatim ≤240 chars from the chunk text). NEVER leave document_name or snippet blank.


Output format (STRICT): Your final response is a single JSON object that matches the response schema exactly. Do NOT wrap it in an extra key like {"<ClassName>": ...} or {"result": ...}. Every required string field is a string, not a nested object. Every required nested model is included with all of its required fields populated. Never omit required fields; never add unspecified ones."""

def _find_placeholders(doc: DocxDocument) -> list[str]:
    seen: list[str] = []
    for para in doc.paragraphs:
        for match in PLACEHOLDER_RE.finditer(para.text):
            name = match.group(1)
            if name not in seen:
                seen.append(name)
    return seen


def _replace(doc: DocxDocument, name: str, value: str) -> None:
    token = "{{" + name + "}}"
    for para in doc.paragraphs:
        if token in para.text:
            for run in para.runs:
                if token in run.text:
                    run.text = run.text.replace(token, value)


async def run(template: Path, question_hint: str) -> None:
    mcp = MCPServerStdio(
        command=os.environ.get("KS_MCP_COMMAND", ".venv/bin/ks-mcp"),
        args=(os.environ.get("KS_MCP_ARGS", "") or "").split(),
        env={
            "KS_API_KEY": os.environ.get("KS_API_KEY", ""),
            "KS_BASE_URL": os.environ.get("KS_BASE_URL", ""),
        },
    )
    agent = Agent(
        model=f"openai:{os.environ.get('MODEL', 'gpt-4o')}",
        mcp_servers=[mcp],
        system_prompt=PROMPT,
        output_type=FieldAnswer,
        retries=4,
        output_retries=4,
    )
    doc = DocxDocument(template)
    fields = _find_placeholders(doc)
    if not fields:
        sys.exit(f"No {{{{placeholders}}}} found in {template}.")
    evidence: dict[str, dict] = {}
    async with agent.run_mcp_servers():
        for field in fields:
            request = f"Field: {field.replace('_', ' ')}\nContext: {question_hint}"
            result = await agent.run(request)
            answer = result.output
            _replace(doc, field, answer.value)
            evidence[field] = answer.model_dump()
    out_docx = template.with_suffix(".filled.docx")
    out_json = template.with_suffix(".citations.json")
    doc.save(out_docx)
    out_json.write_text(json.dumps(evidence, indent=2))
    print(json.dumps({"filled_docx": str(out_docx), "citations_json": str(out_json), "fields_filled": len(evidence)}, indent=2))


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "--template", type=Path, required=True, help="A .docx with {{field_name}} placeholders."
    )
    p.add_argument("--hint", default="Answer from our security & privacy policies.")
    args = p.parse_args()
    if not os.environ.get("KS_API_KEY") or not os.environ.get("OPENAI_API_KEY"):
        sys.exit("Set KS_API_KEY and OPENAI_API_KEY.")
    asyncio.run(run(args.template, args.hint))


if __name__ == "__main__":
    main()
