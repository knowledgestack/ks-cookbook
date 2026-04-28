"""ChiroCRM visit autopilot CLI — one call, three artifacts."""

import argparse
import asyncio
import os
import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from chiro_autopilot.agent import run_autopilot
from chiro_autopilot.schema import VisitAutopilotOutput


def _render_coding_xlsx(out: VisitAutopilotOutput, path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Coding"
    ws.append(
        [
            "patient_id",
            "visit_date",
            "code",
            "type",
            "description",
            "units",
            "fee_cad",
            "fee_source_chunk",
            "justification_chunk",
        ]
    )
    for c in out.codes:
        ws.append(
            [
                out.patient_id,
                out.visit_date,
                c.code,
                c.code_type,
                c.description,
                c.units,
                c.fee_cad,
                f"{c.fee_citation.document_name}:{c.fee_citation.chunk_id}",
                f"{c.justification_citation.document_name}:{c.justification_citation.chunk_id}",
            ]
        )
    wb.save(path)


def _render_prior_auth_docx(out: VisitAutopilotOutput, path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Prior Authorization — Patient {out.patient_id}", level=1)
    doc.add_paragraph(f"Visit date: {out.visit_date}")
    doc.add_paragraph(f"Chief complaint: {out.chief_complaint}")
    if out.prior_auth_letter_body:
        doc.add_paragraph(out.prior_auth_letter_body)
    doc.add_heading("Coverage criteria", level=2)
    for cr in out.prior_auth_criteria:
        status = "MET" if cr.met else "NOT DOCUMENTED"
        doc.add_paragraph(f"[{status}] {cr.criterion}", style="List Bullet")
        doc.add_paragraph(
            f"    Policy: {cr.policy_citation.document_name} (chunk:{cr.policy_citation.chunk_id})"
        )
        if cr.chart_citation:
            doc.add_paragraph(
                f"    Chart: {cr.chart_citation.document_name} (chunk:{cr.chart_citation.chunk_id})"
            )
    doc.save(path)


def _render_patient_plan_md(out: VisitAutopilotOutput, path: Path) -> None:
    lines: list[str] = [
        f"# Your care plan — visit {out.visit_date}",
        "",
        "Your chiropractor recommends the following plan. Each step is",
        "based on our clinic's care protocols.",
        "",
    ]
    for step in out.patient_plan:
        lines.append(f"## {step.visit_window} — {step.focus}")
        lines.append(step.plain_language_instruction)
        c = step.protocol_citation
        lines.append(f"*Source: {c.document_name} (chunk:{c.chunk_id})*")
        lines.append("")
    lines.append("## Notes from your clinician")
    lines.append(out.clinician_note)
    path.write_text("\n".join(lines))


def main() -> None:
    p = argparse.ArgumentParser(description="ChiroCRM visit autopilot — 3 artifacts.")
    p.add_argument("--patient-id", required=True)
    p.add_argument("--visit-date", required=True, help="YYYY-MM-DD")
    p.add_argument(
        "--corpus-folder",
        default=os.environ.get(
            "CHIRO_CORPUS_FOLDER_ID",
            "ab926019-ac7a-579f-bfda-6c52a13c5f41",
        ),
    )
    p.add_argument("--model", default=os.environ.get("MODEL", "gpt-4o"))
    p.add_argument("--out-dir", type=Path, default=Path("flagships/chiro_visit_autopilot"))
    args = p.parse_args()

    if not os.environ.get("KS_API_KEY") or not os.environ.get("OPENAI_API_KEY"):
        sys.exit("Set KS_API_KEY and OPENAI_API_KEY in .env.")

    result = asyncio.run(
        run_autopilot(
            patient_id=args.patient_id,
            visit_date=args.visit_date,
            corpus_folder_id=args.corpus_folder,
            model=args.model,
        )
    )
    args.out_dir.mkdir(parents=True, exist_ok=True)
    coding = args.out_dir / "sample_output_coding.xlsx"
    pa = args.out_dir / "sample_output_prior_auth.docx"
    plan = args.out_dir / "sample_output_patient_plan.md"
    _render_coding_xlsx(result, coding)
    _render_prior_auth_docx(result, pa)
    _render_patient_plan_md(result, plan)
    print(
        f"Wrote 3 artifacts — coding={coding} prior_auth={pa} plan={plan} "
        f"codes={len(result.codes)} pa_required={result.prior_auth_required} "
        f"plan_steps={len(result.patient_plan)}"
    )


if __name__ == "__main__":
    main()
