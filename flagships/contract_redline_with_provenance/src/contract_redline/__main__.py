"""Contract redline CLI — produces a Word memo with traceable edits."""

import argparse
import asyncio
import os
import sys
from pathlib import Path

from docx import Document

from contract_redline.agent import produce_redline
from contract_redline.schema import RedlineMemo


def _render_docx(memo: RedlineMemo, out: Path) -> None:
    doc = Document()
    doc.add_heading(f"Redline memo — {memo.counterparty_document}", level=1)
    doc.add_paragraph(f"Playbook: {memo.playbook_name}")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(memo.summary)
    doc.add_heading("Proposed redlines", level=2)
    for i, r in enumerate(memo.redlines, 1):
        doc.add_heading(f"{i}. {r.clause_title}  [{r.risk_tier.upper()}]", level=3)
        doc.add_paragraph("Original:", style="Intense Quote").add_run(f"\n{r.original_text}")
        doc.add_paragraph("Proposed:", style="Intense Quote").add_run(f"\n{r.proposed_text}")
        doc.add_paragraph(r.rationale)
        doc.add_paragraph(
            f"Offending clause: {r.offending_clause_citation.document_name} "
            f"(chunk:{r.offending_clause_citation.chunk_id})"
        )
        doc.add_paragraph(
            f"Playbook rule: {r.playbook_rule_citation.document_name} "
            f"(chunk:{r.playbook_rule_citation.chunk_id})"
        )
    if memo.acceptable_as_is:
        doc.add_heading("Acceptable as drafted", level=2)
        for note in memo.acceptable_as_is:
            doc.add_paragraph(note, style="List Bullet")
    doc.save(out)


def main() -> None:
    p = argparse.ArgumentParser(
        description="Produce a Word redline memo with dual-corpus citations."
    )
    p.add_argument("--playbook-name", default="firm_playbook")
    p.add_argument("--inbound-name", default="counterparty_draft")
    p.add_argument(
        "--playbook-folder",
        default=os.environ.get(
            "PLAYBOOK_FOLDER_ID",
            "a4bdb206-d45a-50fa-9b62-071966226eb8",
        ),
    )
    p.add_argument(
        "--draft-folder",
        default=os.environ.get(
            "DRAFT_FOLDER_ID",
            "a4bdb206-d45a-50fa-9b62-071966226eb8",
        ),
    )
    p.add_argument("--model", default=os.environ.get("MODEL", "gpt-4o"))
    p.add_argument("--out", type=Path, default=Path("redline-memo.docx"))
    args = p.parse_args()

    if not os.environ.get("KS_API_KEY") or not os.environ.get("OPENAI_API_KEY"):
        sys.exit("Set KS_API_KEY and OPENAI_API_KEY in .env.")

    memo = asyncio.run(
        produce_redline(
            playbook_folder_id=args.playbook_folder,
            draft_folder_id=args.draft_folder,
            inbound_name=args.inbound_name,
            playbook_name=args.playbook_name,
            model=args.model,
        )
    )
    _render_docx(memo, args.out)
    blockers = sum(1 for r in memo.redlines if r.risk_tier == "blocker")
    print(
        f"Wrote {args.out} — redlines={len(memo.redlines)} "
        f"blockers={blockers} acceptable={len(memo.acceptable_as_is)}"
    )


if __name__ == "__main__":
    main()
