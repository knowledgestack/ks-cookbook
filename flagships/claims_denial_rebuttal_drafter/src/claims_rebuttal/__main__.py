"""Claims denial rebuttal CLI."""

import argparse
import asyncio
import os
import sys
from pathlib import Path

from docx import Document

from claims_rebuttal.agent import draft_rebuttal
from claims_rebuttal.schema import RebuttalLetter


def _render_docx(letter: RebuttalLetter, out: Path) -> None:
    doc = Document()
    doc.add_heading(f"Appeal — Denial {letter.denial_code}", level=1)
    doc.add_paragraph(f"Payer: {letter.payer}")
    doc.add_paragraph(f"Patient: {letter.patient_id}")
    doc.add_paragraph(f"Service: {letter.service_in_question}")
    doc.add_paragraph(letter.opening_paragraph)
    doc.add_heading("Coverage criteria", level=2)
    for i, c in enumerate(letter.criteria, 1):
        status = "MET" if c.is_met else "NOT DOCUMENTED"
        doc.add_heading(f"{i}. {c.criterion} [{status}]", level=3)
        doc.add_paragraph(
            f"Policy: {c.policy_citation.document_name} "
            f"(chunk:{c.policy_citation.chunk_id}) — \u201c{c.policy_citation.snippet}\u201d"
        )
        for ev in c.supporting_evidence:
            doc.add_paragraph(
                f"Chart evidence: {ev.document_name} (chunk:{ev.chunk_id}) — "
                f"\u201c{ev.snippet}\u201d",
                style="List Bullet",
            )
    doc.add_heading("Requested action", level=2)
    doc.add_paragraph(letter.requested_action)
    doc.add_paragraph(letter.closing_paragraph)
    doc.save(out)


def main() -> None:
    p = argparse.ArgumentParser(description="Draft a denial rebuttal letter.")
    p.add_argument("--patient-id", required=True)
    p.add_argument("--denial-code", required=True, help="e.g. CO-50, CO-197")
    p.add_argument("--payer", default="BCBS")
    p.add_argument(
        "--service",
        default="Lumbar epidural steroid injection",
        help="Service / CPT description being appealed.",
    )
    p.add_argument(
        "--chart-folder",
        default=os.environ.get(
            "PATIENT_CHART_FOLDER_ID",
            "ab926019-ac7a-579f-bfda-6c52a13c5f41",
        ),
    )
    p.add_argument(
        "--policy-folder",
        default=os.environ.get(
            "PAYER_POLICY_FOLDER_ID",
            "ab926019-ac7a-579f-bfda-6c52a13c5f41",
        ),
    )
    p.add_argument("--model", default=os.environ.get("MODEL", "gpt-4o"))
    p.add_argument("--out", type=Path, default=Path("appeal-letter.docx"))
    args = p.parse_args()

    if not os.environ.get("KS_API_KEY") or not os.environ.get("OPENAI_API_KEY"):
        sys.exit("Set KS_API_KEY and OPENAI_API_KEY in .env.")

    letter = asyncio.run(
        draft_rebuttal(
            patient_id=args.patient_id,
            denial_code=args.denial_code,
            payer=args.payer,
            service=args.service,
            chart_folder_id=args.chart_folder,
            policy_folder_id=args.policy_folder,
            model=args.model,
        )
    )
    _render_docx(letter, args.out)
    met = sum(1 for c in letter.criteria if c.is_met)
    print(
        f"Wrote {args.out} — denial={letter.denial_code} criteria={len(letter.criteria)} met={met}"
    )


if __name__ == "__main__":
    main()
