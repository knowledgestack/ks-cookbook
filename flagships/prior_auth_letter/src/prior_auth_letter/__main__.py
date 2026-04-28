"""CLI entry for the prior-authorization-letter demo."""

import argparse
import asyncio
import os
import sys
from pathlib import Path

from docx import Document as DocxDocument

from prior_auth_letter.agent import draft_letter
from prior_auth_letter.schema import PriorAuthLetter


def _render_docx(letter: PriorAuthLetter, out_path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Prior Authorization / Appeal Letter", level=0)
    doc.add_paragraph(f"Member: {letter.member_name}  |  DOB: {letter.member_dob}")
    doc.add_paragraph(f"Member ID: {letter.member_id}")
    doc.add_paragraph(f"Ordering Provider: {letter.ordering_provider}")
    doc.add_paragraph(f"Requested Service: {letter.requested_service}")

    for heading, body in (
        ("Clinical Scenario", letter.clinical_scenario),
        ("Medical Necessity", letter.medical_necessity),
        ("Prior Therapies", letter.prior_therapies),
        ("Supporting Evidence", letter.supporting_evidence or "None provided."),
    ):
        doc.add_heading(heading, level=1)
        for para in body.split("\n\n"):
            stripped = para.strip()
            if stripped:
                doc.add_paragraph(stripped)

    doc.add_heading("Citations", level=1)
    for idx, c in enumerate(letter.citations, start=1):
        snippet = (c.snippet or "").replace("\n", " ").strip()
        doc.add_paragraph(
            f"[{idx}] {c.document_name} — [chunk:{c.chunk_id}] \u201c{snippet}\u201d"
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _load_request(in_path: Path | None) -> str:
    if in_path is not None:
        return in_path.read_text(encoding="utf-8")
    default = Path(__file__).resolve().parents[2] / "sample_inputs" / "scenario.txt"
    return default.read_text(encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Draft a cited prior-auth letter from a clinical scenario."
    )
    parser.add_argument(
        "--in",
        dest="in_path",
        type=Path,
        default=None,
        help="Scenario text file (default: sample_inputs/scenario.txt).",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("prior-auth-letter.docx"),
        help="Output .docx (default: prior-auth-letter.docx).",
    )
    parser.add_argument(
        "--corpus-folder",
        default=os.environ.get("CORPUS_FOLDER_ID", ""),
        help="Folder id for the health-plan policy corpus.",
    )
    parser.add_argument(
        "--model",
        default=os.environ.get("PRIOR_AUTH_MODEL", "openai:gpt-4o"),
        help="pydantic-ai model id (default: openai:gpt-4o).",
    )
    args = parser.parse_args()

    if not os.environ.get("KS_API_KEY"):
        sys.exit("KS_API_KEY is not set. See the README.")
    if not args.corpus_folder:
        sys.exit("--corpus-folder (or CORPUS_FOLDER_ID env var) is required.")

    request = _load_request(args.in_path)
    letter = asyncio.run(
        draft_letter(
            request,
            corpus_folder_id=args.corpus_folder,
            model=args.model,
        )
    )
    _render_docx(letter, args.out)
    print(
        f"Wrote {args.out} — {len(letter.citations)} citation(s); "
        f"{len(letter.medical_necessity)} chars of medical-necessity prose."
    )


if __name__ == "__main__":
    main()
