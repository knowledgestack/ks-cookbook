"""Render a ``BriefOutput`` to a Word (.docx) file.

Mirrors the pattern used inside the KS backend's ``generate_docx`` agent tool
but stays self-contained — no imports from ks-backend.
"""

from pathlib import Path

from docx import Document as DocxDocument

from research_brief.schema import BriefOutput


def write_docx(brief: BriefOutput, out_path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading(brief.title, level=0)

    for section in brief.sections:
        doc.add_heading(section.heading, level=1)
        for paragraph in section.body.split("\n\n"):
            stripped = paragraph.strip()
            if stripped:
                doc.add_paragraph(stripped)

    doc.add_heading("References", level=1)
    for citation in brief.citations:
        snippet = (citation.snippet or "").replace("\n", " ").strip()
        page = f", p. {citation.page}" if citation.page else ""
        doc.add_paragraph(
            f"[{citation.id}] {citation.document_name}{page}. "
            f"Chunk {citation.chunk_id}. \u201c{snippet}\u201d"
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
