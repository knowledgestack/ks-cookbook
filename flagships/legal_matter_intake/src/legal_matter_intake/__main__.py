"""Sertain legal matter intake CLI."""

import argparse
import asyncio
import os
import sys
from pathlib import Path

from docx import Document

from legal_matter_intake.agent import produce_dossier
from legal_matter_intake.schema import MatterIntakeDossier


def _render_docx(d: MatterIntakeDossier, out: Path) -> None:
    doc = Document()
    doc.add_heading(f"Matter Intake Dossier — {d.client}", level=1)
    doc.add_paragraph(f"Matter: {d.matter}")
    doc.add_paragraph(f"Overall risk: {d.overall_risk.upper()}")

    doc.add_heading("Conflicts analysis", level=2)
    if not d.conflicts_analysis:
        doc.add_paragraph("No conflicts identified against the firm's database.")
    for c in d.conflicts_analysis:
        doc.add_paragraph(
            f"[{c.hit_type.upper()}] {c.adverse_party} — {c.past_matter} → {c.resolution}",
            style="List Bullet",
        )
        doc.add_paragraph(f"    Source: {c.citation.document_name} (chunk:{c.citation.chunk_id})")

    doc.add_heading("Risk factors", level=2)
    for r in d.risk_factors:
        doc.add_paragraph(f"[{r.severity.upper()}] {r.title}", style="List Bullet")
        doc.add_paragraph(f"    {r.narrative}")
        doc.add_paragraph(f"    Source: {r.citation.document_name} (chunk:{r.citation.chunk_id})")

    doc.add_heading("Fee estimate", level=2)
    fe = d.fee_estimate
    doc.add_paragraph(f"Practice area: {fe.practice_area}")
    doc.add_paragraph(f"Staffing: {fe.staffing_model}")
    doc.add_paragraph(f"Estimate: {fe.currency} {fe.low_range:,} – {fe.high_range:,}")
    doc.add_paragraph(f"Assumptions: {fe.assumptions}")
    doc.add_paragraph(f"    Source: {fe.citation.document_name} (chunk:{fe.citation.chunk_id})")

    doc.add_heading("Required ethics disclosures", level=2)
    for e in d.required_disclosures:
        doc.add_paragraph(f"{e.rule_reference}: {e.requirement}", style="List Bullet")
        doc.add_paragraph(f"    Source: {e.citation.document_name} (chunk:{e.citation.chunk_id})")

    doc.add_heading("Recommended engagement terms", level=2)
    for t in d.recommended_engagement_terms:
        doc.add_paragraph(t, style="List Number")

    doc.add_heading("Partner notes", level=2)
    doc.add_paragraph(d.partner_notes)
    doc.save(out)


def main() -> None:
    p = argparse.ArgumentParser(description="Produce a legal matter intake dossier.")
    p.add_argument("--client", required=True)
    p.add_argument("--matter", required=True)
    p.add_argument(
        "--corpus-folder",
        default=os.environ.get(
            "LEGAL_CORPUS_FOLDER_ID",
            "ab926019-ac7a-579f-bfda-6c52a13c5f41",
        ),
    )
    p.add_argument("--model", default=os.environ.get("MODEL", "gpt-4o"))
    p.add_argument("--out", type=Path, default=Path("matter-intake.docx"))
    args = p.parse_args()

    if not os.environ.get("KS_API_KEY") or not os.environ.get("OPENAI_API_KEY"):
        sys.exit("Set KS_API_KEY and OPENAI_API_KEY in .env.")

    dossier = asyncio.run(
        produce_dossier(
            client=args.client,
            matter=args.matter,
            corpus_folder_id=args.corpus_folder,
            model=args.model,
        )
    )
    _render_docx(dossier, args.out)
    print(
        f"Wrote {args.out} — risk={dossier.overall_risk} "
        f"conflicts={len(dossier.conflicts_analysis)} "
        f"risks={len(dossier.risk_factors)} "
        f"disclosures={len(dossier.required_disclosures)}"
    )


if __name__ == "__main__":
    main()
