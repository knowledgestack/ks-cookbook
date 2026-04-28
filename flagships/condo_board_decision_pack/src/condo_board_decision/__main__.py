"""Condo board decision pack CLI — outputs board-ready docx."""

import argparse
import asyncio
import os
import sys
from pathlib import Path

from docx import Document

from condo_board_decision.agent import decide
from condo_board_decision.schema import BoardDecisionPack


def _render_docx(pack: BoardDecisionPack, out: Path) -> None:
    doc = Document()
    doc.add_heading("Board Decision Memorandum", level=1)
    doc.add_paragraph(f"Unit: {pack.unit}").bold = True
    doc.add_paragraph(f"Request: {pack.request_summary}")
    doc.add_paragraph()
    doc.add_heading(f"VERDICT: {pack.verdict}", level=2)
    doc.add_paragraph(f"Vote threshold: {pack.required_vote_threshold}")

    doc.add_heading("Governing rules", level=2)
    for r in pack.governing_rules:
        doc.add_paragraph(f"[{r.source}] {r.rule_summary}", style="List Bullet")
        doc.add_paragraph(
            f"    Source: {r.citation.document_name} "
            f"(chunk:{r.citation.chunk_id}) — \u201c{r.citation.snippet}\u201d"
        )

    if pack.precedents:
        doc.add_heading("Precedents from prior board minutes", level=2)
        for p in pack.precedents:
            doc.add_paragraph(
                f"{p.meeting_date} — {p.similar_request} → {p.outcome}",
                style="List Bullet",
            )
            doc.add_paragraph(
                f"    Source: {p.citation.document_name} (chunk:{p.citation.chunk_id})"
            )

    if pack.reserve_impact:
        doc.add_heading("Reserve-fund impact", level=2)
        doc.add_paragraph(pack.reserve_impact.note)
        if pack.reserve_impact.citation:
            c = pack.reserve_impact.citation
            doc.add_paragraph(f"    Source: {c.document_name} (chunk:{c.chunk_id})")

    if pack.conditions:
        doc.add_heading("Conditions", level=2)
        for cond in pack.conditions:
            doc.add_paragraph(cond, style="List Number")

    doc.add_heading("Rationale", level=2)
    doc.add_paragraph(pack.rationale)

    doc.add_heading("Recommended motion", level=2)
    doc.add_paragraph(pack.recommended_motion)
    doc.save(out)


def main() -> None:
    p = argparse.ArgumentParser(description="Board decision pack for a unit-owner request.")
    p.add_argument("--request", required=True)
    p.add_argument("--unit", default="4B")
    p.add_argument(
        "--corpus-folder",
        default=os.environ.get(
            "CONDO_CORPUS_FOLDER_ID",
            "ab926019-ac7a-579f-bfda-6c52a13c5f41",
        ),
    )
    p.add_argument("--model", default=os.environ.get("MODEL", "gpt-4o"))
    p.add_argument("--out", type=Path, default=Path("board-decision.docx"))
    args = p.parse_args()

    if not os.environ.get("KS_API_KEY") or not os.environ.get("OPENAI_API_KEY"):
        sys.exit("Set KS_API_KEY and OPENAI_API_KEY in .env.")

    pack = asyncio.run(
        decide(
            request=args.request,
            unit=args.unit,
            corpus_folder_id=args.corpus_folder,
            model=args.model,
        )
    )
    _render_docx(pack, args.out)
    print(
        f"Wrote {args.out} — verdict={pack.verdict} "
        f"rules={len(pack.governing_rules)} precedents={len(pack.precedents)} "
        f"conditions={len(pack.conditions)}"
    )


if __name__ == "__main__":
    main()
